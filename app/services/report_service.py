import os
import json
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from reportlab.pdfgen import canvas
from jinja2 import Template
from migrations.migrations import get_db_connection
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import logging

logger = logging.getLogger(__name__)

class ReportService:
    def __init__(self):
        self.report_dir = "reports"
        if not os.path.exists(self.report_dir):
            os.makedirs(self.report_dir)

    def generate_method_rating(self, time_interval):
        end_date = datetime.now()
        start_date = end_date - timedelta(hours=time_interval)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        query = '''
            SELECT 
                http_method || ' ' || api_path as method,
                COUNT(*) as call_count,
                MAX(created_at) as last_call
            FROM LogsRequests
            WHERE created_at BETWEEN ? AND ?
            GROUP BY http_method, api_path
            ORDER BY call_count DESC
        '''
        
        cursor.execute(query, (start_date.strftime('%Y-%m-%d %H:%M:%S'), 
                             end_date.strftime('%Y-%m-%d %H:%M:%S')))
        
        method_stats = []
        for row in cursor.fetchall():
            method_stats.append({
                'method': row[0],
                'call_count': row[1],
                'last_call': row[2]
            })
            
        conn.close()
        return method_stats

    def generate_entity_rating(self, time_interval):
        end_date = datetime.now()
        start_date = end_date - timedelta(hours=time_interval)
        
        conn = get_db_connection()
        cursor = conn.cursor()

        # Проверяем все записи в таблице без фильтров
        cursor.execute("SELECT * FROM ChangeLogs LIMIT 5")
        sample_data = cursor.fetchall()
        if sample_data:
            columns = [description[0] for description in cursor.description]
            formatted_data = [dict(zip(columns, row)) for row in sample_data]
            logger.info(f"Sample data from ChangeLogs: {formatted_data}")
        
        # Проверяем временной диапазон
        cursor.execute("SELECT MIN(created_at), MAX(created_at) FROM ChangeLogs")
        time_range = cursor.fetchone()
        logger.info(f"ChangeLogs time range: from {time_range[0]} to {time_range[1]}")
        
        # Проверяем формат даты в базе
        cursor.execute("SELECT created_at FROM ChangeLogs LIMIT 1")
        date_sample = cursor.fetchone()
        if date_sample:
            logger.info(f"Sample date format in ChangeLogs: {date_sample[0]}")
        
        query = '''
            SELECT 
                COALESCE(entity_type, 'Unknown') as entity_type,
                COUNT(*) as change_count,
                MAX(created_at) as last_change
            FROM ChangeLogs
            WHERE datetime(created_at) BETWEEN datetime(?) AND datetime(?)
            GROUP BY entity_type
            ORDER BY change_count DESC
        '''
        
        logger.info(f"Executing query with dates: {start_date} - {end_date}")
        cursor.execute(query, (start_date.strftime('%Y-%m-%d %H:%M:%S'), 
                             end_date.strftime('%Y-%m-%d %H:%M:%S')))
        
        results = cursor.fetchall()
        logger.info(f"Query results: {results}")
        
        entity_stats = []
        for row in results:
            entity_stats.append({
                'entity_type': row[0],
                'change_count': row[1],
                'last_change': row[2]
            })
            
        conn.close()
        return entity_stats

    def generate_user_rating(self, time_interval):
        end_date = datetime.now()
        start_date = end_date - timedelta(hours=time_interval)
        
        conn = get_db_connection()
        cursor = conn.cursor()

        # Проверяем данные в LogsRequests
        cursor.execute("SELECT * FROM LogsRequests LIMIT 5")
        sample_data = cursor.fetchall()
        if sample_data:
            columns = [description[0] for description in cursor.description]
            formatted_data = [dict(zip(columns, row)) for row in sample_data]
            logger.info(f"Sample data from LogsRequests: {formatted_data}")
        
        # Проверяем временной диапазон
        cursor.execute("SELECT MIN(created_at), MAX(created_at) FROM LogsRequests")
        time_range = cursor.fetchone()
        logger.info(f"LogsRequests time range: from {time_range[0]} to {time_range[1]}")
        
        # Проверяем формат даты в базе
        cursor.execute("SELECT created_at FROM LogsRequests LIMIT 1")
        date_sample = cursor.fetchone()
        if date_sample:
            logger.info(f"Sample date format in LogsRequests: {date_sample[0]}")
        
        # Запросы и авторизации пользователей
        query_requests = '''
            SELECT 
                COALESCE(user_id, 'anonymous') as user_id,
                COUNT(*) as request_count,
                SUM(CASE WHEN api_path LIKE '/api/auth%' THEN 1 ELSE 0 END) as auth_count,
                MAX(created_at) as last_activity
            FROM LogsRequests
            WHERE datetime(created_at) BETWEEN datetime(?) AND datetime(?)
            GROUP BY user_id
        '''
        
        logger.info(f"Executing requests query with dates: {start_date} - {end_date}")
        cursor.execute(query_requests, (start_date.strftime('%Y-%m-%d %H:%M:%S'), 
                                      end_date.strftime('%Y-%m-%d %H:%M:%S')))
        
        results = cursor.fetchall()
        if results:
            formatted_results = [{'user_id': row[0], 'request_count': row[1], 
                                'auth_count': row[2], 'last_activity': row[3]} 
                               for row in results]
            logger.info(f"Requests query results: {formatted_results}")
        
        user_stats = {}
        for row in results:
            user_id = row[0]
            user_stats[user_id] = {
                'user_id': user_id,
                'request_count': row[1],
                'auth_count': row[2],
                'change_count': 0,
                'last_activity': row[3]
            }
        
        # Изменения пользователей
        query_changes = '''
            SELECT 
                COALESCE(created_by, 'anonymous') as user_id,
                COUNT(*) as change_count,
                MAX(created_at) as last_change
            FROM ChangeLogs
            WHERE datetime(created_at) BETWEEN datetime(?) AND datetime(?)
            GROUP BY created_by
        '''
        
        logger.info(f"Executing changes query with dates: {start_date} - {end_date}")
        cursor.execute(query_changes, (start_date.strftime('%Y-%m-%d %H:%M:%S'), 
                                     end_date.strftime('%Y-%m-%d %H:%M:%S')))
        
        results = cursor.fetchall()
        if results:
            formatted_results = [{'user_id': row[0], 'change_count': row[1], 
                                'last_change': row[2]} for row in results]
            logger.info(f"Changes query results: {formatted_results}")
        
        for row in results:
            user_id = row[0]
            if user_id not in user_stats:
                user_stats[user_id] = {
                    'user_id': user_id,
                    'request_count': 0,
                    'auth_count': 0,
                    'change_count': row[1],
                    'last_activity': row[2]
                }
            else:
                user_stats[user_id]['change_count'] = row[1]
                if row[2] > user_stats[user_id]['last_activity']:
                    user_stats[user_id]['last_activity'] = row[2]
        
        conn.close()
        return sorted(list(user_stats.values()), 
                     key=lambda x: (x['request_count'] + x['change_count']), 
                     reverse=True)

    def generate_report(self, time_interval):
        end_date = datetime.now()
        start_date = end_date - timedelta(hours=time_interval)
        
        report_data = {
            'report_info': {
                'type': 'System Usage Statistics',
                'start_date': start_date.isoformat(),
                'end_date': end_date.isoformat(),
                'generated_at': datetime.now().isoformat()
            },
            'method_statistics': self.generate_method_rating(time_interval),
            'entity_statistics': self.generate_entity_rating(time_interval),
            'user_statistics': self.generate_user_rating(time_interval)
        }
        
        return report_data

    def save_report(self, report, format='json'):
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"report_{timestamp}.{format}"
        filepath = os.path.join(self.report_dir, filename)

        if format == 'json':
            with open(filepath, 'w') as f:
                json.dump(report, f, indent=2)
        elif format in ['xlsx', 'csv']:
            df_methods = pd.DataFrame(report['method_statistics'])
            df_entities = pd.DataFrame(report['entity_statistics'])
            df_users = pd.DataFrame(report['user_statistics'])

            if format == 'xlsx':
                with pd.ExcelWriter(filepath) as writer:
                    df_methods.to_excel(writer, sheet_name='Methods Rating', index=False)
                    df_entities.to_excel(writer, sheet_name='Entities Rating', index=False)
                    df_users.to_excel(writer, sheet_name='Users Rating', index=False)
            else:  # csv
                df_methods.to_csv(filepath, index=False)

        elif format == 'docx':
            doc = Document()
            doc.add_heading('System Usage Statistics Report', 0)
            doc.add_paragraph(f"Report Period: {report['report_info']['start_date']} - {report['report_info']['end_date']}")
            doc.add_paragraph(f"Generated at: {report['report_info']['generated_at']}")
            
            doc.add_heading('Methods Statistics', level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'Method'
            table.rows[0].cells[1].text = 'Call Count'
            table.rows[0].cells[2].text = 'Last Call'
            
            for method in report['method_statistics']:
                row = table.add_row()
                row.cells[0].text = method['method']
                row.cells[1].text = str(method['call_count'])
                row.cells[2].text = method['last_call']
            
            doc.save(filepath)

        return filepath

    def send_report(self, filepath):
        smtp_server = os.getenv('SMTP_SERVER')
        smtp_port = int(os.getenv('SMTP_PORT'))
        smtp_username = os.getenv('SMTP_USERNAME')
        smtp_password = os.getenv('SMTP_PASSWORD')
        admin_emails = os.getenv('ADMIN_EMAILS').split(',')

        msg = MIMEMultipart()
        msg['Subject'] = 'Activity Report'
        msg['From'] = smtp_username
        msg['To'] = ', '.join(admin_emails)

        with open(filepath, 'rb') as f:
            attachment = MIMEApplication(f.read(), _subtype=os.path.splitext(filepath)[1][1:])
            attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(filepath))
            msg.attach(attachment)

        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_username, smtp_password)
            server.send_message(msg)

        os.remove(filepath)  # Удаляем файл после отправки 